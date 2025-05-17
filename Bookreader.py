# Set Tesseract paths directly in code - BEFORE importing pytesseract
import os
import sys
os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'
os.environ['PATH'] += r';C:\Program Files\Tesseract-OCR'

# Now import pytesseract
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Test if Tesseract is working before proceeding
try:
    print("Testing Tesseract installation...")
    from PIL import Image
    import numpy as np
    test_img = np.zeros((100, 100, 3), dtype=np.uint8)
    test_img.fill(255)  # White image
    test_result = pytesseract.image_to_string(Image.fromarray(test_img))
    print("Tesseract test successful!")
except Exception as e:
    print(f"Tesseract test failed: {str(e)}")
    # Show error in GUI
    import tkinter as tk
    from tkinter import messagebox
    root = tk.Tk()
    root.withdraw()
    messagebox.showerror("Tesseract Error", 
                         f"Tesseract OCR is not properly configured: {str(e)}\n\n"
                         "Please check your Tesseract installation and try again.")
    sys.exit(1)

# Rest of imports
import cv2
import numpy as np
import pyttsx3
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from datetime import datetime
from docx import Document
from langdetect import detect
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import threading
import queue
from PIL import Image, ImageTk

# --- CONFIGURATION ---
PREVIEW_WIDTH, PREVIEW_HEIGHT = 800, 600

model_name = "google/flan-t5-base"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
pipe = pipeline("text2text-generation", model=model, tokenizer=tokenizer)

# Initialize document
doc = Document()
ocr_queue = queue.Queue(maxsize=1)
ai_queue = queue.Queue(maxsize=1)
TEXT_BUFFER = []
EXPLANATION_CACHE = {}

def list_available_cameras(max_cameras=5):
    available = []
    for i in range(max_cameras):
        try:
            cap = cv2.VideoCapture(i, cv2.CAP_DSHOW)
            if cap.read()[0]:
                available.append(i)
            cap.release()
        except:
            pass
    return available or [0]

# FIXED speak function to avoid "run loop already started" error
def speak(text):
    try:
        # Use a completely new engine instance without threading
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()
        engine.stop()
    except Exception as e:
        print(f"TTS Error: {str(e)}")

def extract_text(frame):
    try:
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        text = pytesseract.image_to_string(gray).strip()
        if text:
            try:
                lang = detect(text)
                text = pytesseract.image_to_string(gray, lang=lang).strip()
            except:
                pass  # Fall back to default text if language detection fails
        return text
    except Exception as e:
        return f"OCR Error: {str(e)}"

def ocr_worker(frame_queue, ocr_queue):
    while True:
        frame = frame_queue.get()
        if frame is None:
            break
        text = extract_text(frame)
        try:
            ocr_queue.put_nowait(text)
        except queue.Full:
            pass

def ai_worker(ai_task_queue, ai_result_queue):
    while True:
        text = ai_task_queue.get()
        if text is None:
            break
        if text in EXPLANATION_CACHE:
            explanation = EXPLANATION_CACHE[text]
        else:
            try:
                prompt = f"Explain in simple terms: {text}"
                explanation = pipe(prompt, max_new_tokens=100)[0]['generated_text']
                EXPLANATION_CACHE[text] = explanation
            except Exception as e:
                explanation = f"AI Error: {str(e)}"
        try:
            ai_result_queue.put_nowait(explanation)
        except queue.Full:
            pass

class SmartBookGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Smart Book Assistant")
        self.geometry("1500x900")
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.iconbitmap("icon.ico") if os.path.exists("icon.ico") else None

        # Camera
        self.cap = None
        self.running = False
        self.paused = False
        self.frame_count = 0
        self.frame_queue = queue.Queue(maxsize=1)
        self.ai_task_queue = queue.Queue(maxsize=1)
        self.ai_result_queue = queue.Queue(maxsize=1)
        self.current_camera = 0
        self.cam_url = None

        # Layout
        self.create_widgets()
        self.start_workers()
        self.start_camera()
        self.after(100, self.update_camera)
        self.after(200, self.update_ocr_and_ai)

    def create_widgets(self):
        # Left: Camera preview and camera controls
        left_frame = tk.Frame(self)
        left_frame.pack(side=tk.LEFT, padx=10, pady=10, anchor="n")

        self.video_label = tk.Label(left_frame, width=PREVIEW_WIDTH, height=PREVIEW_HEIGHT, bg="black")
        self.video_label.pack()

        # Camera controls
        control_frame = tk.Frame(left_frame)
        control_frame.pack(fill=tk.Y, padx=5, pady=5)

        tk.Label(control_frame, text="Select Camera:").pack()
        self.camera_var = tk.StringVar()
        self.cam_selector = ttk.Combobox(control_frame, textvariable=self.camera_var, state="readonly")
        self.cam_selector['values'] = list_available_cameras()
        self.cam_selector.current(0)
        self.cam_selector.pack()
        self.cam_selector.bind("<<ComboboxSelected>>", self.change_camera)

        # Pause/Resume button
        self.pause_btn = ttk.Button(control_frame, text="Pause", command=self.toggle_pause)
        self.pause_btn.pack(pady=5)

        # IP camera entry
        tk.Label(control_frame, text="IP Camera:").pack(pady=(10,0))
        ip_frame = tk.Frame(control_frame)
        ip_frame.pack(fill=tk.X, pady=5)
        self.ip_entry = ttk.Entry(ip_frame)
        self.ip_entry.pack(side=tk.LEFT, padx=(0,5), fill=tk.X, expand=True)
        ttk.Button(ip_frame, text="Connect", command=self.connect_ip_cam).pack(side=tk.RIGHT)

        # Right: Text, explanation, notes, save
        right_frame = tk.Frame(self)
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        tk.Label(right_frame, text="Scanned Text Preview", font=("Helvetica", 12, "bold")).pack()
        self.text_preview = scrolledtext.ScrolledText(right_frame, wrap=tk.WORD, height=8, font=("Consolas", 12))
        self.text_preview.pack(pady=5, fill=tk.BOTH, expand=True)

        tk.Label(right_frame, text="AI Explanation", font=("Helvetica", 12, "bold")).pack()
        self.explanation_preview = scrolledtext.ScrolledText(right_frame, wrap=tk.WORD, height=6, font=("Consolas", 12))
        self.explanation_preview.pack(pady=5, fill=tk.BOTH, expand=True)

        # Notes
        notes_frame = tk.LabelFrame(right_frame, text="Your Notes (auto-filled with scans)")
        notes_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        self.notes_input = tk.Text(notes_frame, height=10, font=("Consolas", 12))
        self.notes_input.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
        ttk.Button(notes_frame, text="Save Notes", command=self.save_notes).pack(pady=5)

        # Save & Exit
        ttk.Button(right_frame, text="Save & Exit", command=self.save_and_exit).pack(pady=10)

    def start_workers(self):
        threading.Thread(target=ocr_worker, args=(self.frame_queue, ocr_queue), daemon=True).start()
        threading.Thread(target=ai_worker, args=(self.ai_task_queue, self.ai_result_queue), daemon=True).start()

    def start_camera(self):
        self.running = False
        if self.cap:
            self.cap.release()
        try:
            if hasattr(self, 'cam_url') and self.cam_url:
                self.cap = cv2.VideoCapture(self.cam_url)
            else:
                self.current_camera = int(self.cam_selector.get())
                self.cap = cv2.VideoCapture(self.current_camera)
            
            if not self.cap.isOpened():
                messagebox.showerror("Camera Error", "Could not open selected camera")
                return
                
            self.running = True
        except Exception as e:
            messagebox.showerror("Camera Error", f"Error: {str(e)}")

    def update_camera(self):
        if self.running and self.cap and self.cap.isOpened():
            if not self.paused:
                self.frame_count += 1
                ret, frame = self.cap.read()
                if ret:
                    # Show preview (large)
                    img = cv2.cvtColor(cv2.resize(frame, (PREVIEW_WIDTH, PREVIEW_HEIGHT)), cv2.COLOR_BGR2RGB)
                    imgtk = ImageTk.PhotoImage(Image.fromarray(img))
                    self.video_label.imgtk = imgtk
                    self.video_label.configure(image=imgtk)
                    # Send frame for OCR every 3rd frame
                    if self.frame_count % 3 == 0 and self.frame_queue.empty():
                        self.frame_queue.put(frame.copy())
        self.after(50, self.update_camera)

    def update_ocr_and_ai(self):
        # OCR result
        try:
            text = ocr_queue.get_nowait()
            self.text_preview.delete(1.0, tk.END)
            self.text_preview.insert(tk.END, text)
            if text and text not in TEXT_BUFFER and not text.startswith("OCR Error"):
                TEXT_BUFFER.append(text)
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                doc.add_paragraph(f"[Captured Notes] - {timestamp}")
                doc.add_paragraph(f"Text: {text}")
                doc.add_paragraph("\n")
                # --- Add scanned text to notes box live ---
                self.notes_input.insert(tk.END, f"\n[Scan @ {timestamp}]:\n{text}\n")
                self.notes_input.see(tk.END)  # Auto-scroll to bottom
                
                # AI explanation - only for substantial text that's not an error
                if len(text.split()) > 5 and self.ai_task_queue.empty():
                    self.ai_task_queue.put(text)
        except queue.Empty:
            pass

        # AI explanation result
        try:
            explanation = self.ai_result_queue.get_nowait()
            self.explanation_preview.delete(1.0, tk.END)
            self.explanation_preview.insert(tk.END, explanation)
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"[Explanation] - {timestamp}")
            doc.add_paragraph(explanation)
            doc.add_paragraph("\n")
            
            # Use a separate thread for TTS to avoid blocking the UI
            threading.Thread(target=lambda: speak(explanation), daemon=True).start()
        except queue.Empty:
            pass

        self.after(200, self.update_ocr_and_ai)

    def toggle_pause(self):
        self.paused = not self.paused
        self.pause_btn.config(text="Resume" if self.paused else "Pause")

    def change_camera(self, event=None):
        if hasattr(self, 'cam_url'):
            delattr(self, 'cam_url')
        self.start_camera()

    def connect_ip_cam(self):
        ip = self.ip_entry.get().strip()
        if ip:
            self.cam_url = f'http://{ip}:4747/mjpegfeed'
            self.start_camera()

    def save_notes(self):
        notes = self.notes_input.get("1.0", tk.END).strip()
        if notes:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"[User Notes] - {timestamp}")
            doc.add_paragraph(notes)
            self.notes_input.delete("1.0", tk.END)
            messagebox.showinfo("Notes Saved", "Your notes have been saved to the document.")

    def save_and_exit(self):
        self.running = False
        if self.cap:
            self.cap.release()
        self.save_notes()
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            title="Save your notes as...",
            initialfile="SmartBookNotes.docx"
        )
        if file_path:
            doc.save(file_path)
            # Use a simple message instead of TTS for exit
            messagebox.showinfo("Done", f"All notes saved to {file_path}")
        else:
            messagebox.showinfo("Cancelled", "Save operation was cancelled.")
        self.destroy()

    def on_closing(self):
        self.running = False
        if self.cap:
            self.cap.release()
        self.destroy()

if __name__ == "__main__":
    doc.add_heading('Smart Book Notes', 0)
    SmartBookGUI().mainloop()
